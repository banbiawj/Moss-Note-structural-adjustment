from __future__ import annotations

import html
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from bs4 import BeautifulSoup
from fastapi import UploadFile


@dataclass
class ParsedDocument:
    filename: str
    text: str
    html: str


async def parse_upload_file(file: UploadFile) -> ParsedDocument:
    raw = await file.read()
    filename = file.filename or "document.txt"
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ""}:
        text = raw.decode("utf-8", errors="ignore")
        content_html = _plain_text_to_html(text)
    elif suffix in {".md", ".markdown"}:
        text = raw.decode("utf-8", errors="ignore")
        content_html = _markdown_to_html(text)
    elif suffix == ".docx":
        text, content_html = _docx_to_html(raw)
    elif suffix == ".pdf":
        text = _pdf_to_text(raw)
        content_html = _plain_text_to_html(text)
    else:
        raise ValueError("仅支持 .txt、.md、.markdown、.docx、.pdf 文件")

    return ParsedDocument(
        filename=filename,
        text=_clean_text(text),
        html=ensure_block_ids(content_html),
    )


def ensure_block_ids(content_html: str) -> str:
    soup = BeautifulSoup(f"<main>{content_html}</main>", "html.parser")
    main = soup.find("main")
    if main is None:
        return content_html

    block_tags = {"div", "p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "blockquote"}
    for child in list(main.find_all(recursive=False)):
        if not getattr(child, "name", None) or child.name not in block_tags:
            continue
        if child.name == "div":
            child["id"] = child.get("id") or _block_id()
            continue
        wrapper = soup.new_tag("div", id=_block_id())
        child.wrap(wrapper)

    return "".join(str(child) for child in main.contents)


def _plain_text_to_html(text: str) -> str:
    blocks = []
    for paragraph in _split_paragraphs(text):
        blocks.append(f"<p>{html.escape(paragraph)}</p>")
    return "\n".join(blocks) or '<p id="empty-document">空文档</p>'


def _markdown_to_html(text: str) -> str:
    import markdown

    return markdown.markdown(text, extensions=["extra", "sane_lists"])


def _docx_to_html(raw: bytes) -> tuple[str, str]:
    from docx import Document

    document = Document(BytesIO(raw))
    text_parts: list[str] = []
    html_parts: list[str] = []

    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if not content:
            continue
        text_parts.append(content)
        style_name = (paragraph.style.name or "").lower()
        escaped = html.escape(content)
        if "heading 1" in style_name:
            html_parts.append(f"<h1>{escaped}</h1>")
        elif "heading 2" in style_name:
            html_parts.append(f"<h2>{escaped}</h2>")
        elif "heading 3" in style_name:
            html_parts.append(f"<h3>{escaped}</h3>")
        else:
            html_parts.append(f"<p>{escaped}</p>")

    return "\n".join(text_parts), "\n".join(html_parts)


def _pdf_to_text(raw: bytes) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(BytesIO(raw)) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n\n".join(chunk for chunk in chunks if chunk.strip())


def _split_paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    return [part.strip() for part in normalized.split("\n\n") if part.strip()]


def _clean_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(lines).strip()


def _block_id() -> str:
    return f"moss-block-{uuid4().hex[:12]}"

